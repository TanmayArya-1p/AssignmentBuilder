import docx
import yaml 
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import os
from pygments import highlight
from pygments.lexers import PythonLexer
from pygments.formatters import HtmlFormatter
from tqdm import tqdm
import json

build_config = json.load(open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json"), 'r'))

max_lines_per_page = build_config["MAX_LINES_PER_PAGE"]


co =  open("abconfig.yaml", 'r')
config = yaml.safe_load(co)

doc = docx.Document()
p = doc.add_paragraph()

p.paragraph_format.line_spacing = 1
p.paragraph_format.space_after = 0
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = p.add_run(config["title"])

run.bold = True
run.font.name = build_config["PRIMARY_FONT"]

run.font.size = docx.shared.Pt(16)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(config["byline"])
run.font.name = build_config["PRIMARY_FONT"]
run.font.size = docx.shared.Pt(10)

doc.add_paragraph()
print("Creating DOCX...")

linesused = 3
prob = config["problems"]
for i in tqdm(prob):
    if linesused >= max_lines_per_page:
        doc.add_page_break()
        linesused = 0

    p = doc.add_paragraph()
    p.add_run(i.upper()+":").font.name = build_config["PRIMARY_FONT"]
    linesused += 1
    p = doc.add_paragraph()
    r = p.add_run("Source Code:")
    r.font.name = build_config["PRIMARY_FONT"]
    linesused += 1

    r.font.bold = True
    p = doc.add_paragraph()
    with open(prob[i]["source"] , "r") as file:
        src = file.read()
        r = p.add_run(src)
        linesused += src.count("\n")

        r.font.name = build_config["CODE_FONT"]
        subprocess.Popen(f"gcc {prob[i]['source']}")
        command = f"{os.getcwd()}\\a.exe"

        process = subprocess.Popen(
        command,
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
        )
        stdout, stderr = process.communicate(input=prob[i]["testcase"])
        
        console_content_before = (
            f"{os.getcwd()}> {command}\n"
        )
        console_content_after = (
            f"{stdout}\n"
            f"{stderr}"
        )

        if linesused >= max_lines_per_page:
            doc.add_page_break()
            linesused = 0

        p = doc.add_paragraph()
        r = p.add_run("Output:")
        r.font.bold = True
        r.font.name = build_config["PRIMARY_FONT"]
        p = doc.add_paragraph()
        r = p.add_run(console_content_before)
        r.font.size = docx.shared.Pt(10)
        r.font.name =build_config["CODE_FONT"]


        r = p.add_run(prob[i]["testcase"])
        r.font.size = docx.shared.Pt(10)
        r.font.name = build_config["CODE_FONT"]
        r.font.italic = True

        r = p.add_run(console_content_after)
        r.font.size = docx.shared.Pt(10)
        r.font.name = build_config["CODE_FONT"]
        linesused += console_content_after.count("\n")+1 + console_content_before.count("\n")+1

import webbrowser
import docx2pdf


doc.save(f"{config['title']}.docx")
print("Generating PDF...")
docx2pdf.convert(f"{config['title']}.docx", f"{config['title']}.pdf")
webbrowser.open(f"{config['title']}.pdf")